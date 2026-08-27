#!/usr/bin/env python3
"""Assemble PIGSTI final report package: heatmap figure + Excel results + Word report."""

from __future__ import annotations

import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, Cm

ROOT = Path(__file__).resolve().parent.parent
DOCS = ROOT / "docs"
FINAL = ROOT / "final"

HEATMAP_PNG = DOCS / "pathogen_detection_scores_heatmap.png"
HEATMAP_PDF = DOCS / "pathogen_detection_scores_heatmap.pdf"
SRC_PATHOGEN_XLSX = FINAL / "pathogen_summary_all_samples.xlsx"
SRC_HOST_XLSX = FINAL / "host_mtdna_summary_all_samples.xlsx"

OUT_RESULTS_XLSX = FINAL / "pathogen_results.xlsx"
OUT_REPORT_DOCX = FINAL / "PIGSTI_final_report.docx"


def _read_excel(path: Path) -> dict[str, pd.DataFrame]:
    if not path.exists():
        return {}
    xl = pd.ExcelFile(path)
    return {name: pd.read_excel(path, sheet_name=name) for name in xl.sheet_names}


def build_results_excel() -> Path:
    """Write a clean multi-sheet Excel results workbook for the final report."""
    FINAL.mkdir(parents=True, exist_ok=True)
    pathogen_sheets = _read_excel(SRC_PATHOGEN_XLSX)
    host_sheets = _read_excel(SRC_HOST_XLSX)

    # Prefer first sheet; normalise name
    pathogen_df = next(iter(pathogen_sheets.values())) if pathogen_sheets else pd.DataFrame()
    host_df = next(iter(host_sheets.values())) if host_sheets else pd.DataFrame()

    with pd.ExcelWriter(OUT_RESULTS_XLSX, engine="openpyxl") as writer:
        if not pathogen_df.empty:
            pathogen_df.to_excel(writer, sheet_name="pathogen_summary", index=False)
        if not host_df.empty:
            host_df.to_excel(writer, sheet_name="host_mtdna", index=False)

        # Compact sample × pathogen score-like view when columns exist
        if not pathogen_df.empty and {"Sample", "Pathogen"}.issubset(pathogen_df.columns):
            value_col = None
            for cand in ("Score", "Escore", "Krakenuniq_reads", "Coverage"):
                if cand in pathogen_df.columns:
                    value_col = cand
                    break
            if value_col is not None:
                wide = pathogen_df.pivot_table(
                    index="Pathogen",
                    columns="Sample",
                    values=value_col,
                    aggfunc="first",
                )
                wide.to_excel(writer, sheet_name=f"matrix_{value_col}")

        readme = pd.DataFrame(
            {
                "Item": [
                    "pathogen_summary",
                    "host_mtdna",
                    "heatmap_png",
                    "heatmap_pdf",
                    "report_docx",
                ],
                "Description": [
                    "Per-sample pathogen metrics (KrakenUniq, E-value, mapping, damage, ANI, …)",
                    "Host / mtDNA cohort summary",
                    "Pathogen detection scores heatmap figure (PNG)",
                    "Pathogen detection scores heatmap figure (PDF)",
                    "Word report embedding the heatmap and tables",
                ],
                "Path": [
                    str(OUT_RESULTS_XLSX.name) + " :: pathogen_summary",
                    str(OUT_RESULTS_XLSX.name) + " :: host_mtdna",
                    "pathogen_detection_scores_heatmap.png",
                    "pathogen_detection_scores_heatmap.pdf",
                    OUT_REPORT_DOCX.name,
                ],
            }
        )
        readme.to_excel(writer, sheet_name="README", index=False)

    # Also keep canonical summary filename in sync with a named sheet
    if not pathogen_df.empty:
        with pd.ExcelWriter(SRC_PATHOGEN_XLSX, engine="openpyxl") as writer:
            pathogen_df.to_excel(writer, sheet_name="pathogen_summary", index=False)

    print(f"Wrote {OUT_RESULTS_XLSX}")
    return OUT_RESULTS_XLSX


def _add_df_table(doc: Document, df: pd.DataFrame, max_rows: int = 40) -> None:
    if df.empty:
        doc.add_paragraph("(No data)")
        return
    show = df.head(max_rows).fillna("")
    cols = [str(c) for c in show.columns]
    table = doc.add_table(rows=1 + len(show), cols=len(cols))
    table.style = "Table Grid"
    for j, c in enumerate(cols):
        table.rows[0].cells[j].text = c
    for i, (_, row) in enumerate(show.iterrows(), start=1):
        for j, c in enumerate(cols):
            val = row[show.columns[j]]
            table.rows[i].cells[j].text = "" if pd.isna(val) else str(val)
    if len(df) > max_rows:
        doc.add_paragraph(f"… showing first {max_rows} of {len(df)} rows (see Excel for full table).")


def build_report_docx(pathogen_df: pd.DataFrame, host_df: pd.DataFrame) -> Path:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("PIGSTI Final Results Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Cohort-level pathogen detection heatmap and Excel results tables "
        "for the PIGSTI publication package."
    )

    # Heatmap
    doc.add_heading("1. Pathogen detection scores heatmap", level=1)
    doc.add_paragraph(
        "Heatmap of multi-criteria pathogen detection scores across samples "
        "(attached figure). Cell labels show criteria passed / criteria evaluated."
    )
    if HEATMAP_PNG.exists():
        doc.add_picture(str(HEATMAP_PNG), width=Inches(6.5))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph(f"(Missing figure: {HEATMAP_PNG})")
    doc.add_paragraph(
        f"Source files: {HEATMAP_PNG.name} · {HEATMAP_PDF.name}"
    ).runs[0].italic = True

    # Excel results
    doc.add_heading("2. Excel results", level=1)
    doc.add_paragraph(
        f"Full numeric tables are in {OUT_RESULTS_XLSX.name} "
        f"(also {SRC_PATHOGEN_XLSX.name} and {SRC_HOST_XLSX.name})."
    )

    doc.add_heading("2.1 Pathogen summary", level=2)
    _add_df_table(doc, pathogen_df)

    doc.add_heading("2.2 Host / mtDNA summary", level=2)
    _add_df_table(doc, host_df)

    doc.add_heading("3. File checklist", level=1)
    for item in (
        f"{HEATMAP_PNG.name} — detection-score heatmap (PNG)",
        f"{HEATMAP_PDF.name} — detection-score heatmap (PDF)",
        f"{OUT_RESULTS_XLSX.name} — combined Excel results workbook",
        f"{SRC_PATHOGEN_XLSX.name} — pathogen summary Excel",
        f"{SRC_HOST_XLSX.name} — host/mtDNA summary Excel",
        f"{OUT_REPORT_DOCX.name} — this Word report",
    ):
        doc.add_paragraph(item, style="List Bullet")

    doc.save(OUT_REPORT_DOCX)
    print(f"Wrote {OUT_REPORT_DOCX}")
    return OUT_REPORT_DOCX


def main() -> int:
    FINAL.mkdir(parents=True, exist_ok=True)

    # Paste attached heatmap into final/
    for src in (HEATMAP_PNG, HEATMAP_PDF):
        if not src.exists():
            print(f"WARNING: missing {src}")
            continue
        dst = FINAL / src.name
        shutil.copy2(src, dst)
        print(f"Copied {src.name} -> {dst}")

    build_results_excel()

    pathogen_sheets = _read_excel(SRC_PATHOGEN_XLSX)
    host_sheets = _read_excel(SRC_HOST_XLSX)
    pathogen_df = next(iter(pathogen_sheets.values())) if pathogen_sheets else pd.DataFrame()
    host_df = next(iter(host_sheets.values())) if host_sheets else pd.DataFrame()
    build_report_docx(pathogen_df, host_df)

    print("\nFinal package contents:")
    for p in sorted(FINAL.iterdir()):
        print(f"  {p.name}  ({p.stat().st_size} bytes)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
