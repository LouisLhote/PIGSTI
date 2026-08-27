#!/usr/bin/env python3
"""Build METRICS_DEFINITIONS.docx for Google Docs / Word."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_label_paragraph(doc: Document, label: str, body: str) -> None:
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    run_label.bold = True
    p.add_run(body)


def add_bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "auto")
        borders.append(element)
    tbl_pr.append(borders)


def build_doc(output_path: Path) -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("PIGSTI Pathogen Authentication Metrics", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Definitions of all metrics used in the PIGSTI multi-criteria pathogen "
        "detection framework, suitable for a Results or Methods section."
    )

    # 1
    add_heading(doc, "1. KrakenUniq number of reads (Krakenuniq_reads)", 1)
    add_label_paragraph(
        doc,
        "Definition",
        "The number of sequencing reads assigned to a taxon at the clade level in the "
        "KrakenUniq report (column 'reads'), encompassing reads assigned to the taxon and "
        "all its taxonomic descendants.",
    )
    add_label_paragraph(
        doc,
        "Interpretation",
        "A minimum read count filters weak or sporadic k-mer hits. Default threshold: "
        "≥ 50 reads (per-pathogen overrides possible via Pathogen_spreadsheet.csv).",
    )
    add_label_paragraph(
        doc,
        "Source",
        "Breitwieser, F. P. et al. KrakenUniq: confident and fast metagenomics "
        "classification using unique k-mer counts. Genome Biology 19, 198 (2018).",
    )

    # 2
    add_heading(doc, "2. E-value (Guellil_et_al_Evalue)", 1)
    add_label_paragraph(
        doc,
        "Definition",
        "A KrakenUniq-based confidence score for distinguishing true from false taxonomic assignments:",
    )
    doc.add_paragraph("E = (K / R) × C", style="Intense Quote")
    doc.add_paragraph(
        "where K = number of unique k-mers assigned to the taxon, R = number of assigned reads, "
        "and C = estimated k-mer coverage (genome breadth) from the KrakenUniq report."
    )
    add_label_paragraph(
        doc,
        "Interpretation",
        "Higher values indicate reads spread across more of the genome with good k-mer uniqueness "
        "relative to read count. Low E with many reads often suggests repetitive or misclassified "
        "k-mers. Default detection threshold: E > 0.001.",
    )
    add_label_paragraph(
        doc,
        "Source",
        "Guellil, M. et al. Ancient herpes simplex 1 genomes reveal recent viral structure in "
        "Eurasia. Science Advances 8, eabo4435 (2022). See also Guellil, M. et al. An 8000 years "
        "old genome reveals the Neolithic origin of the zoonosis Brucella melitensis. "
        "Nature Communications 15, 5781 (2024).",
    )

    # 3
    add_heading(doc, "3. Relative entropy (Relative_entropy)", 1)
    add_label_paragraph(
        doc,
        "Definition",
        "Shannon entropy of the distribution of read alignment start positions along the reference "
        "genome, computed in non-overlapping windows (100 bp and 1,000 bp), normalized by the "
        "maximum entropy attainable if the same number of reads were evenly distributed across all windows:",
    )
    doc.add_paragraph(
        "H = −Σ p_i log₂(p_i)",
        style="Intense Quote",
    )
    doc.add_paragraph(
        "H_rel = H / H_max = [−Σ p_i log₂(p_i)] / log₂(N_windows)",
        style="Intense Quote",
    )
    doc.add_paragraph("where:")
    add_bullet(
        doc,
        "H = raw Shannon entropy of the read-start distribution across windows. It measures how "
        "spread out mapped read start positions are along the genome: low H means reads cluster in a "
        "few windows (uneven coverage); high H means reads are distributed more evenly.",
    )
    add_bullet(
        doc,
        "p_i = fraction of read starts in window i (reads in window i / total reads)",
    )
    add_bullet(
        doc,
        "H_max = log₂(N_windows), the maximum entropy if all reads were perfectly evenly distributed "
        "across all windows",
    )
    add_bullet(
        doc,
        "H_rel = relative entropy (0–1 scale), the value reported by PIGSTI",
    )
    add_label_paragraph(
        doc,
        "Interpretation",
        "Values range from 0 to 1. Higher values indicate reads distributed evenly along the genome "
        "(expected for genuine pathogen signal). Low values suggest reads clustered in few regions "
        "(contamination, mis-mapping, or fragmented reference). Default thresholds: ≥ 0.9 for "
        "bacteria/archaea; ≥ 0.7 for viruses and genomes < 10 kb.",
    )
    add_label_paragraph(
        doc,
        "Source",
        "Sikora, M. et al. The spatiotemporal distribution of human pathogens in ancient Eurasia. "
        "Nature 643, 1011–1019 (2025). Implemented following the pathopipe workflow.",
    )

    # 4
    add_heading(doc, "4. Edit distance 1 — Damaged reads (Edit_distance_decay_quality)", 1)
    add_label_paragraph(
        doc,
        "Definition",
        "Decay Quality Score (0–1) for the damage-enriched read subset. Reads are classified as "
        "'damaged' if they show terminal deamination (5′ C→T and/or 3′ G→A within the first/last "
        "5 aligned bases). The score quantifies whether the edit-distance (NM tag) distribution "
        "follows the expected monotonic decay pattern (most reads at NM = 0, declining with "
        "increasing mismatches). For damaged reads, bins starting at edit distance ≥ 1 are used.",
    )
    doc.add_paragraph("The composite score integrates:")
    add_bullet(doc, "Monotonicity of the decay (35% weight)")
    add_bullet(doc, "Dominance ratio: count at NM=0 / count at NM=1 (25%)")
    add_bullet(doc, "Fraction of reads at NM=0 (15%)")
    add_bullet(doc, "Peak position penalty (−10% if peak ≠ 0)")
    add_bullet(doc, "Exponential decay rate (10%)")
    add_bullet(doc, "R² of exponential fit (5%)")
    add_label_paragraph(
        doc,
        "Interpretation",
        "Higher scores indicate a clean descending mismatch profile consistent with authentic "
        "ancient DNA. Default pass threshold: ≥ 0.65.",
    )
    add_label_paragraph(
        doc,
        "Source",
        "Inspired by HOPS edit-distance authentication: Hübler, R. et al. HOPS: automated "
        "detection and authentication of pathogen DNA in archaeological remains. "
        "Genome Biology 20, 280 (2019).",
    )

    # 5
    add_heading(doc, "5. Edit distance 2 — Default reads (Edit_distance_decay_quality_default)", 1)
    add_label_paragraph(
        doc,
        "Definition",
        "The same Decay Quality Score computed on the non-damaged (default) read subset — mapped "
        "reads without terminal deamination patterns in the 5′/3′ windows.",
    )
    add_label_paragraph(
        doc,
        "Interpretation",
        "In authentic ancient samples, damaged reads should show a stronger descending "
        "edit-distance pattern than non-damaged reads. Comparing the two scores helps distinguish "
        "true ancient signal from modern contamination. Default pass threshold: ≥ 0.55.",
    )
    add_label_paragraph(doc, "Source", "Same as Edit distance 1; HOPS-style damage/default split.")

    # 6
    add_heading(doc, "6. Damage (Damage_5p_CtoT)", 1)
    add_label_paragraph(
        doc,
        "Definition",
        "Frequency of cytosine-to-thymine substitutions at the first position of the 5′ end of "
        "aligned reads, reflecting characteristic post-mortem deamination of ancient DNA.",
    )
    add_label_paragraph(
        doc,
        "Interpretation",
        "Elevated 5′ C→T rates support ancient origin of the mapped reads. Default pass threshold: "
        "≥ 0.01 (1%).",
    )
    add_label_paragraph(
        doc,
        "Source",
        "DamageProfiler: Neukamm, J. et al. Bioinformatics 37, 3652–3653 (2021). "
        "Damage model: Jónsson, H. et al. Bioinformatics 29, 1682–1684 (2013).",
    )

    # 7
    add_heading(doc, "7. ANI (ANI)", 1)
    add_label_paragraph(
        doc,
        "Definition",
        "Average Nucleotide Identity — the mean fraction of bases in mapped reads matching the "
        "reference, expressed as a percentage:",
    )
    doc.add_paragraph(
        "ANI ≈ (1 − mismatches / bases mapped) × 100",
        style="Intense Quote",
    )
    doc.add_paragraph("Computed from samtools stats output fields (mismatches and bases mapped).")
    add_label_paragraph(
        doc,
        "Interpretation",
        "Higher ANI indicates closer sequence similarity to the reference genome. Default pass "
        "threshold: > 96.5% (bacteria); > 95% for viruses.",
    )
    add_label_paragraph(
        doc,
        "Source",
        "Sikora, M. et al. Nature 643, 1011–1019 (2025); Li, H. et al. Bioinformatics 25, "
        "2078–2079 (2009).",
    )

    # 8
    add_heading(doc, "8. Breadth ratio (Breadth_ratio)", 1)
    add_label_paragraph(
        doc,
        "Definition",
        "Ratio of observed to expected breadth of genomic coverage:",
    )
    doc.add_paragraph(
        "Breadth ratio = B_observed / B_expected",
        style="Intense Quote",
    )
    doc.add_paragraph("where:")
    add_bullet(doc, "B_observed = fraction of reference positions covered by ≥ 1 read")
    add_bullet(
        doc,
        "B_expected = 1 − e^(−mean depth)  (Poisson expectation for uniform coverage)",
    )
    add_label_paragraph(
        doc,
        "Interpretation",
        "Values near 1.0 indicate even coverage relative to read depth. Values < 1 suggest patchy "
        "or clustered mapping. Default pass threshold: ≥ 0.8.",
    )
    add_label_paragraph(doc, "Source", "Sikora, M. et al. Nature 643, 1011–1019 (2025).")

    # 9
    add_heading(doc, "9. Mapping ratio (Read_mapping_ratio)", 1)
    add_label_paragraph(
        doc,
        "Definition",
        "Consistency metric between KrakenUniq classification and reference mapping:",
    )
    doc.add_paragraph(
        "Mapping ratio = Mapped reads (BWA/Bowtie2) / KrakenUniq reads",
        style="Intense Quote",
    )
    add_label_paragraph(doc, "Interpretation", "")
    add_bullet(doc, "≈ 1.0: Good agreement between k-mer classification and alignment.")
    add_bullet(doc, "> 1.0: More reads map than KrakenUniq assigned.")
    add_bullet(doc, "< 1.0: Fewer reads map than classified.")
    doc.add_paragraph("Default pass threshold: ≥ 0.5.")
    add_label_paragraph(doc, "Source", "PIGSTI pipeline cross-validation metric.")

    # 10
    add_heading(doc, "10. Genus ranking (Genus_ranking)", 1)
    add_label_paragraph(
        doc,
        "Definition",
        "Rank position of the target pathogen species among all species in the same genus detected "
        "in the KrakenUniq report, ordered by descending clade read count. Reported as #1, #2, etc.",
    )
    add_label_paragraph(
        doc,
        "Interpretation",
        "Rank #1 indicates the classified species is the dominant hit within its genus, reducing "
        "ambiguity from closely related taxa. PIGSTI awards a detection point when genus ranking = 1.",
    )
    add_label_paragraph(
        doc,
        "Source",
        "PIGSTI implementation following genus-level disambiguation logic (cf. Sikora et al. Nature 2025).",
    )

    # Composite score
    add_heading(doc, "Composite detection score", 1)
    doc.add_paragraph(
        "Each sample–pathogen combination is evaluated against all applicable criteria. Results are "
        "reported as a fraction (e.g. '7/12') representing the number of criteria passed out of the "
        "maximum possible. The maximum depends on whether HOPS is enabled (+3 criteria) and whether "
        "damage-split edit-distance metrics are available (+1 additional criterion)."
    )

    add_heading(doc, "Default thresholds summary", 2)
    headers = ["Metric", "Threshold", "Direction"]
    rows = [
        ("KrakenUniq reads", "≥ 50", "Higher is better"),
        ("Guellil E-value", "> 0.001", "Higher is better"),
        ("Relative entropy (1000 bp)", "≥ 0.9 (bacteria) / ≥ 0.7 (virus)", "Higher is better"),
        ("Edit distance decay (damaged)", "≥ 0.65", "Higher is better"),
        ("Edit distance decay (default)", "≥ 0.55", "Higher is better"),
        ("Damage 5′ C→T", "≥ 0.01", "Higher is better"),
        ("ANI", "> 96.5%", "Higher is better"),
        ("Breadth ratio", "≥ 0.8", "Higher is better"),
        ("Mapping ratio", "≥ 0.5", "Higher is better"),
        ("Genus ranking", "= 1", "Lower is better"),
    ]
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"
    set_table_borders(table)
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            table.rows[i].cells[j].text = val
    table.autofit = True

    # References
    add_heading(doc, "Key references", 1)
    refs = [
        "Breitwieser, F. P., Baker, D. N. & Salzberg, S. L. KrakenUniq: confident and fast metagenomics classification using unique k-mer counts. Genome Biology 19, 198 (2018).",
        "Guellil, M. et al. Ancient herpes simplex 1 genomes reveal recent viral structure in Eurasia. Science Advances 8, eabo4435 (2022).",
        "Guellil, M. et al. An 8000 years old genome reveals the Neolithic origin of the zoonosis Brucella melitensis. Nature Communications 15, 5781 (2024).",
        "Sikora, M. et al. The spatiotemporal distribution of human pathogens in ancient Eurasia. Nature 643, 1011–1019 (2025).",
        "Hübler, R. et al. HOPS: automated detection and authentication of pathogen DNA in archaeological remains. Genome Biology 20, 280 (2019).",
        "Neukamm, J. et al. DamageProfiler: fast damage pattern calculation for ancient DNA. Bioinformatics 37, 3652–3653 (2021).",
        "Li, H. et al. The Sequence Alignment/Map format and SAMtools. Bioinformatics 25, 2078–2079 (2009).",
    ]
    for i, ref in enumerate(refs, start=1):
        doc.add_paragraph(f"{i}. {ref}")

    doc.save(output_path)
    print(f"Wrote {output_path}")


if __name__ == "__main__":
    out = Path(__file__).resolve().parent.parent / "docs" / "METRICS_DEFINITIONS.docx"
    build_doc(out)
